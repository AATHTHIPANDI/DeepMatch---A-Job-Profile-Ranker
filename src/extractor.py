import io
import logging
import fitz  # PyMuPDF
from docx import Document

logger = logging.getLogger(__name__)

def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """Extracts text from PDF bytes using PyMuPDF."""
    text_content = []
    try:
        # Open PDF from bytes stream
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        for page in doc:
            text_content.append(page.get_text())
        doc.close()
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise ValueError(f"Failed to parse PDF file: {e}")
        
    return "\n".join(text_content)

def extract_text_from_docx(docx_bytes: bytes) -> str:
    """Extracts text from DOCX bytes using python-docx."""
    try:
        doc = Document(io.BytesIO(docx_bytes))
        paragraphs_text = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Also extract table cells if any
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        table_text.append(cell.text.strip())
                        
        full_text = paragraphs_text + table_text
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise ValueError(f"Failed to parse DOCX file: {e}")
        
    return "\n".join(full_text)

def extract_text_from_txt(txt_bytes: bytes) -> str:
    """Extracts text from TXT bytes, trying UTF-8 and Latin-1 fallbacks."""
    try:
        return txt_bytes.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return txt_bytes.decode("latin-1")
        except Exception as e:
            logger.error(f"Error decoding TXT file: {e}")
            raise ValueError(f"Failed to parse TXT file: {e}")

def extract_text_from_file(file_bytes: bytes, file_name: str) -> str:
    """
    Orchestrates text extraction based on file extension.
    """
    ext = file_name.split(".")[-1].lower() if "." in file_name else ""
    
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in ["docx", "doc"]:
        return extract_text_from_docx(file_bytes)
    elif ext in ["txt", "md"]:
        return extract_text_from_txt(file_bytes)
    else:
        # Default fallback to plain text decoding
        logger.warning(f"Unknown file extension '.{ext}'. Attempting text decoding.")
        return extract_text_from_txt(file_bytes)
